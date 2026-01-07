import os
from typing import List
import glob
import joblib
import numpy as np
import logging
from .embedder import Embedder

logger = logging.getLogger(__name__)

def extract_pdf_text(file_path: str) -> str:
    """Extract text from a PDF file for indexing."""
    try:
        import PyPDF2
        text_parts = []
        with open(file_path, 'rb') as f:
            reader = PyPDF2.PdfReader(f)
            for page_num, page in enumerate(reader.pages, start=1):
                page_text = page.extract_text() or ''
                for line in page_text.split('\n'):
                    text_parts.append(f"[Page {page_num}] {line}")
        return '\n'.join(text_parts)
    except ImportError:
        return ""  # Skip PDF if PyPDF2 not installed
    except Exception:
        return ""  # Skip unreadable PDFs


def extract_docx_text(file_path: str) -> str:
    """Extract text from a Word document (.docx) for indexing."""
    try:
        from docx import Document
        text_parts = []
        doc = Document(file_path)
        
        # Extract text from paragraphs
        para_num = 0
        for para in doc.paragraphs:
            if para.text.strip():
                para_num += 1
                text_parts.append(f"[Para {para_num}] {para.text}")
        
        # Extract text from tables
        table_num = 0
        for table in doc.tables:
            table_num += 1
            for row in table.rows:
                row_text = ' | '.join(cell.text.strip() for cell in row.cells if cell.text.strip())
                if row_text:
                    text_parts.append(f"[Table {table_num}] {row_text}")
        
        return '\n'.join(text_parts)
    except ImportError:
        logger.warning("python-docx not installed. Install with: pip install python-docx")
        return ""  # Skip DOCX if python-docx not installed
    except Exception as e:
        logger.warning(f"Error reading Word document {file_path}: {e}")
        return ""  # Skip unreadable DOCX files

def load_text_files(data_dir: str, patterns: List[str] = None) -> List[dict]:
    """Load all text-based files from data directory.
    
    Args:
        data_dir: Directory to search for files
        patterns: List of glob patterns to match (default: common text file extensions)
    """
    if patterns is None:
        # Default to common text-based file extensions (including Word docs)
        patterns = ["*.txt", "*.md", "*.csv", "*.json", "*.xml", "*.log", 
                    "*.py", "*.java", "*.js", "*.ts", "*.html", "*.css",
                    "*.yaml", "*.yml", "*.ini", "*.cfg", "*.conf", "*.err", "*.out",
                    "*.sql", "*.sh", "*.bat", "*.ps1", "*.c", "*.cpp", "*.h", 
                    "*.pdf", "*.docx", "*.doc"]
    
    files = []
    for pattern in patterns:
        files.extend(glob.glob(os.path.join(data_dir, "**", pattern), recursive=True))
    
    # Remove duplicates while preserving order
    files = list(dict.fromkeys(files))
    
    docs = []
    for p in files:
        try:
            # Handle PDF files specially
            if p.lower().endswith('.pdf'):
                text = extract_pdf_text(p)
            # Handle Word documents
            elif p.lower().endswith(('.docx', '.doc')):
                text = extract_docx_text(p)
            else:
                with open(p, "r", encoding="utf-8", errors="ignore") as f:
                    text = f.read().strip()
            if text:  # Only include non-empty files
                docs.append({"id": p, "text": text})
        except Exception as e:
            logger.debug(f"Skipping file {p}: {e}")
            pass
    return docs

def build_index(data_dir: str, index_path: str, model_name: str = "tfidf"):
    """Build search index from files in data directory.
    
    Args:
        data_dir: Directory containing files to index
        index_path: Path to save the index file
        model_name: Embedding model name (default: tfidf)
    
    Returns:
        dict: Index containing docs, vectors, model_name, and vectorizer
    """
    logger.info(f"Building index from {data_dir}")
    docs = load_text_files(data_dir)
    texts = [d["text"] for d in docs]
    if len(texts) == 0:
        raise RuntimeError(f"No text files found in {data_dir}")

    logger.info(f"Found {len(docs)} documents, creating embeddings...")
    emb = Embedder(model_name)
    emb.fit(texts)
    vectors = emb.embed(texts)

    index = {
        "docs": docs,
        "vectors": vectors,
        "model_name": model_name,
        "vectorizer": emb.vectorizer,
    }
    os.makedirs(os.path.dirname(index_path), exist_ok=True)
    joblib.dump(index, index_path)
    logger.info(f"Index saved to {index_path}")
    return index


def load_index(index_path: str) -> dict:
    """Load index from file.
    
    Args:
        index_path: Path to the index file
        
    Returns:
        dict: Loaded index
    """
    if not os.path.exists(index_path):
        raise FileNotFoundError(f"Index file not found: {index_path}")
    logger.debug(f"Loading index from {index_path}")
    return joblib.load(index_path)
