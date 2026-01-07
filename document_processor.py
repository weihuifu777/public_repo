"""Document processing utilities for the RAG system."""
import re
from pathlib import Path
from typing import List, Optional
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
import config


class DocumentProcessor:
    """Handles document loading, preprocessing, and chunking."""
    
    def __init__(self, chunk_size: int = config.CHUNK_SIZE, 
                 chunk_overlap: int = config.CHUNK_OVERLAP):
        """Initialize the document processor.
        
        Args:
            chunk_size: Size of text chunks
            chunk_overlap: Overlap between chunks
        """
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
        self.text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size,
            chunk_overlap=chunk_overlap,
            length_function=len,
            separators=["\n\n", "\n", " ", ""]
        )
    
    def load_text_file(self, file_path: Path) -> str:
        """Load content from a text file.
        
        Args:
            file_path: Path to the text file
            
        Returns:
            Content of the file
        """
        with open(file_path, 'r', encoding='utf-8') as f:
            return f.read()
    
    def load_document(self, file_path: Path) -> Optional[str]:
        """Load a document from various file formats.
        
        Args:
            file_path: Path to the document
            
        Returns:
            Content of the document or None if format not supported
        """
        suffix = file_path.suffix.lower()
        
        if suffix == '.txt':
            return self.load_text_file(file_path)
        elif suffix == '.md':
            return self.load_text_file(file_path)
        elif suffix == '.pdf':
            try:
                from pypdf import PdfReader
                reader = PdfReader(str(file_path))
                text = ""
                for page in reader.pages:
                    text += page.extract_text() + "\n"
                return text
            except ImportError:
                print("pypdf not installed. Install with: pip install pypdf")
                return None
        elif suffix == '.docx':
            try:
                from docx import Document as DocxDocument
                doc = DocxDocument(str(file_path))
                return "\n".join([para.text for para in doc.paragraphs])
            except ImportError:
                print("python-docx not installed. Install with: pip install python-docx")
                return None
        else:
            print(f"Unsupported file format: {suffix}")
            return None
    
    def preprocess_text(self, text: str) -> str:
        """Clean and preprocess text.
        
        Args:
            text: Raw text
            
        Returns:
            Preprocessed text
        """
        # Remove excessive whitespace
        text = re.sub(r'\s+', ' ', text)
        # Remove special characters but keep punctuation
        text = re.sub(r'[\x00-\x08\x0b-\x0c\x0e-\x1f\x7f-\x9f]', '', text)
        return text.strip()
    
    def chunk_document(self, text: str, metadata: dict = None) -> List[Document]:
        """Split text into chunks.
        
        Args:
            text: Text to chunk
            metadata: Optional metadata to attach to chunks
            
        Returns:
            List of Document objects
        """
        if metadata is None:
            metadata = {}
        
        # Preprocess text
        text = self.preprocess_text(text)
        
        # Split into chunks
        chunks = self.text_splitter.split_text(text)
        
        # Create Document objects
        documents = []
        for i, chunk in enumerate(chunks):
            doc_metadata = metadata.copy()
            doc_metadata['chunk_id'] = i
            documents.append(Document(page_content=chunk, metadata=doc_metadata))
        
        return documents
    
    def process_file(self, file_path: Path) -> List[Document]:
        """Load and process a file into chunks.
        
        Args:
            file_path: Path to the file
            
        Returns:
            List of Document chunks
        """
        content = self.load_document(file_path)
        if content is None:
            return []
        
        metadata = {
            'source': str(file_path),
            'filename': file_path.name
        }
        
        return self.chunk_document(content, metadata)
    
    def process_directory(self, directory: Path) -> List[Document]:
        """Process all supported files in a directory.
        
        Args:
            directory: Path to the directory
            
        Returns:
            List of all Document chunks
        """
        all_documents = []
        supported_extensions = {'.txt', '.md', '.pdf', '.docx'}
        
        for file_path in directory.rglob('*'):
            if file_path.is_file() and file_path.suffix.lower() in supported_extensions:
                print(f"Processing: {file_path.name}")
                documents = self.process_file(file_path)
                all_documents.extend(documents)
                print(f"  Created {len(documents)} chunks")
        
        return all_documents
